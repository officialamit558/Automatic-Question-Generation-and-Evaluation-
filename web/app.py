from flask import Flask, request, render_template, redirect, url_for, send_file
from werkzeug.utils import secure_filename
from scripts.data_ingestion import parse_uploaded_files
from scripts.question_generation import generate_questions
from scripts.answer_evaluation import evaluate_answer
from scripts.scoring import calculate_score
from web.forms import UploadForm, QuestionForm, EvaluateForm, AnswerForm
import os
import io
from fpdf import FPDF
from docx import Document

app = Flask(__name__)
app.config['SECRET_KEY'] = 'your_secret_key'
app.config['UPLOAD_FOLDER'] = 'C:/Users/amit7/AQGS/data/input/'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # Max upload size: 16MB

# Ensure upload folder exists
if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])

# Define allowed file extensions
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def save_questions_to_pdf(mcqs, short_desc, long_desc, case_studies, filename):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    
    pdf.cell(200, 10, txt="MCQs", ln=True, align='C')
    for i, q in enumerate(mcqs, 1):
        pdf.cell(200, 10, txt=f"{i}. {q}", ln=True)
    
    pdf.cell(200, 10, txt="Short Description Questions", ln=True, align='C')
    for i, q in enumerate(short_desc, 1):
        pdf.cell(200, 10, txt=f"{i}. {q}", ln=True)
    
    pdf.cell(200, 10, txt="Long Description Questions", ln=True, align='C')
    for i, q in enumerate(long_desc, 1):
        pdf.cell(200, 10, txt=f"{i}. {q}", ln=True)
    
    pdf.cell(200, 10, txt="Case Studies", ln=True, align='C')
    for i, q in enumerate(case_studies, 1):
        pdf.cell(200, 10, txt=f"{i}. {q}", ln=True)
    
    pdf.output(filename)

def save_questions_to_docx(mcqs, short_desc, long_desc, case_studies, filename):
    doc = Document()
    
    doc.add_heading('MCQs', level=1)
    for i, q in enumerate(mcqs, 1):
        doc.add_paragraph(f"{i}. {q}")
    
    doc.add_heading('Short Description Questions', level=1)
    for i, q in enumerate(short_desc, 1):
        doc.add_paragraph(f"{i}. {q}")
    
    doc.add_heading('Long Description Questions', level=1)
    for i, q in enumerate(long_desc, 1):
        doc.add_paragraph(f"{i}. {q}")
    
    doc.add_heading('Case Studies', level=1)
    for i, q in enumerate(case_studies, 1):
        doc.add_paragraph(f"{i}. {q}")
    
    doc.save(filename)

@app.route('/')
def index():
    upload_form = UploadForm()
    question_form = QuestionForm()
    evaluate_form = EvaluateForm()
    return render_template('index.html', upload_form=upload_form, question_form=question_form, evaluate_form=evaluate_form)

@app.route('/upload', methods=['POST'])
def upload():
    form = UploadForm()
    if form.validate_on_submit():
        file = form.file.data
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)
            file_type = filename.split('.')[-1]
            content = parse_uploaded_files(file_path, file_type)
            return redirect(url_for('index'))
    return render_template('index.html', upload_form=form, question_form=QuestionForm(), evaluate_form=EvaluateForm(), error="Upload failed.")

@app.route('/generate-questions', methods=['POST'])
def generate():
    form = QuestionForm()
    if form.validate_on_submit():
        content = form.content.data
        num_mcqs = form.num_mcqs.data
        num_short_desc = form.num_short_desc.data
        num_long_desc = form.num_long_desc.data
        num_case_studies = form.num_case_studies.data
        mcqs, short_desc, long_desc, case_studies = generate_questions(content, num_mcqs, num_short_desc, num_long_desc, num_case_studies)
        return render_template('questions.html', mcqs=mcqs, short_desc=short_desc, long_desc=long_desc, case_studies=case_studies)
    return render_template('index.html', upload_form=UploadForm(), question_form=form, evaluate_form=EvaluateForm(), error="Question generation failed.")

@app.route('/download/<format>')
def download(format):
    mcqs = request.args.getlist('mcqs')
    short_desc = request.args.getlist('short_desc')
    long_desc = request.args.getlist('long_desc')
    case_studies = request.args.getlist('case_studies')
    filename = f"generated_questions.{format}"
    if format == 'pdf':
        save_questions_to_pdf(mcqs, short_desc, long_desc, case_studies, filename)
    elif format == 'docx':
        save_questions_to_docx(mcqs, short_desc, long_desc, case_studies, filename)
    return send_file(filename, as_attachment=True)

@app.route('/answer-questions', methods=['POST'])
def answer():
    form = AnswerForm()
    if form.validate_on_submit():
        student_answers = {
            'mcqs': form.mcq_answers.data,
            'short_desc': form.short_desc_answers.data,
            'long_desc': form.long_desc_answers.data,
            'case_studies': form.case_studies_answers.data
        }
        # Process the student answers and evaluate them
        return redirect(url_for('index'))
    return render_template('index.html', upload_form=UploadForm(), question_form=QuestionForm(), evaluate_form=EvaluateForm(), answer_form=form, error="Answer submission failed.")

@app.route('/evaluate-answers', methods=['POST'])
def evaluate():
    form = EvaluateForm()
    if form.validate_on_submit():
        student_answers = form.student_answer.data.split('\n')
        expert_answers = form.expert_answer.data.split('\n')
        scores = calculate_score(student_answers, expert_answers)
        return render_template('scores.html', scores=scores)
    return render_template('index.html', upload_form=UploadForm(), question_form=QuestionForm(), evaluate_form=form, error="Evaluation failed.")

if __name__ == '__main__':
    app.run(debug=True)
